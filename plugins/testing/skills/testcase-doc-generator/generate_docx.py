#!/usr/bin/env python3
"""
Generate DOCX (Word) document from Markdown test cases.

This script converts markdown-formatted test cases into a professionally
formatted Microsoft Word document (.docx) with proper styling, tables,
and structure.

Usage:
    python3 generate_docx.py --input <markdown_file> --output <docx_file> --title <document_title>

Example:
    python3 generate_docx.py \
        --input testcases-user-auth.md \
        --output testcases-user-auth.docx \
        --title "Test Cases: User Authentication"

Requirements:
    - python-docx: Install with `pip install python-docx`
"""

import sys
import argparse
import re
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    print("Error: python-docx library not found.", file=sys.stderr)
    print("Install it with: pip install python-docx", file=sys.stderr)
    sys.exit(1)


def parse_arguments():
    """Parse command-line arguments."""
    parser = argparse.ArgumentParser(
        description="Convert Markdown test cases to DOCX format"
    )
    parser.add_argument(
        "--input",
        "-i",
        required=True,
        help="Input markdown file path"
    )
    parser.add_argument(
        "--output",
        "-o",
        required=True,
        help="Output DOCX file path"
    )
    parser.add_argument(
        "--title",
        "-t",
        required=True,
        help="Document title"
    )
    return parser.parse_args()


def setup_document_styles(doc):
    """Configure document styles for headings and paragraphs."""
    # Title style
    title_style = doc.styles['Title']
    title_font = title_style.font
    title_font.name = 'Calibri'
    title_font.size = Pt(24)
    title_font.bold = True
    title_font.color.rgb = RGBColor(0, 51, 102)  # Dark blue

    # Heading styles
    for i in range(1, 4):
        heading_style = doc.styles[f'Heading {i}']
        heading_font = heading_style.font
        heading_font.name = 'Calibri'
        heading_font.color.rgb = RGBColor(0, 51, 102)  # Dark blue

    # Normal text
    normal_style = doc.styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Calibri'
    normal_font.size = Pt(11)


def parse_markdown_table(lines):
    """Parse markdown table and return rows."""
    rows = []
    for line in lines:
        line = line.strip()
        if line.startswith('|') and line.endswith('|'):
            # Remove leading/trailing pipes and split
            cells = [cell.strip() for cell in line[1:-1].split('|')]
            # Skip separator rows (containing only dashes and pipes)
            if not all(set(cell.strip()) <= {'-', ' ', ':'} for cell in cells):
                rows.append(cells)
    return rows


def add_table_to_doc(doc, rows):
    """Add a table to the document."""
    if not rows or len(rows) < 2:
        return

    # Create table
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = 'Light Grid Accent 1'

    # Populate table
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = cell_text

            # Make header row bold
            if i == 0:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
                # Header background color
                cell._element.get_or_add_tcPr().append(
                    cell._element._new_tag('w:shd')
                )
                cell._element.tcPr.shd.set(
                    '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill',
                    '0070C0'  # Blue
                )


def convert_markdown_to_docx(input_file, output_file, title):
    """Convert markdown file to DOCX with proper formatting."""
    # Read markdown content
    with open(input_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    # Create document
    doc = Document()
    setup_document_styles(doc)

    # Add title
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Process markdown content
    i = 0
    in_code_block = False
    in_table = False
    table_lines = []

    while i < len(lines):
        line = lines[i].rstrip()

        # Skip frontmatter
        if i == 0 and line.startswith('---'):
            # Skip until next ---
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('---'):
                i += 1
            i += 1
            continue

        # Code blocks
        if line.startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_para = doc.add_paragraph()
                code_para.style = 'No Spacing'
            else:
                in_code_block = False
            i += 1
            continue

        if in_code_block:
            code_para = doc.add_paragraph(line)
            code_para.style = 'No Spacing'
            for run in code_para.runs:
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
            i += 1
            continue

        # Tables
        if line.startswith('|'):
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(line)
            i += 1
            continue
        elif in_table:
            # End of table
            in_table = False
            rows = parse_markdown_table(table_lines)
            add_table_to_doc(doc, rows)
            doc.add_paragraph()  # Spacing after table
            table_lines = []

        # Headings
        if line.startswith('#'):
            match = re.match(r'^(#{1,6})\s+(.+)$', line)
            if match:
                level = len(match.group(1))
                heading_text = match.group(2)
                doc.add_heading(heading_text, level=min(level, 3))
                i += 1
                continue

        # Horizontal rules
        if line.strip() in ['---', '***', '___']:
            doc.add_paragraph('_' * 80)
            i += 1
            continue

        # Bold text
        if '**' in line:
            para = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = para.add_run(part[2:-2])
                    run.font.bold = True
                else:
                    para.add_run(part)
            i += 1
            continue

        # Lists
        if re.match(r'^\s*[-*+]\s+', line):
            # Bullet list
            text = re.sub(r'^\s*[-*+]\s+', '', line)
            doc.add_paragraph(text, style='List Bullet')
            i += 1
            continue

        if re.match(r'^\s*\d+\.\s+', line):
            # Numbered list
            text = re.sub(r'^\s*\d+\.\s+', '', line)
            doc.add_paragraph(text, style='List Number')
            i += 1
            continue

        # Regular paragraphs
        if line.strip():
            doc.add_paragraph(line)
        else:
            # Empty line - add spacing
            doc.add_paragraph()

        i += 1

    # Handle any remaining table
    if in_table and table_lines:
        rows = parse_markdown_table(table_lines)
        add_table_to_doc(doc, rows)

    # Save document
    doc.save(output_file)
    print(f"✓ DOCX document created: {output_file}", file=sys.stderr)
    print(f"  Pages: Approximately {len(doc.paragraphs) // 40 + 1}", file=sys.stderr)
    print(f"  Sections: {len([p for p in doc.paragraphs if p.style.name.startswith('Heading')])}", file=sys.stderr)


def main():
    """Main entry point."""
    args = parse_arguments()

    # Validate input file exists
    input_path = Path(args.input)
    if not input_path.exists():
        print(f"Error: Input file not found: {args.input}", file=sys.stderr)
        sys.exit(1)

    # Convert
    try:
        convert_markdown_to_docx(args.input, args.output, args.title)
    except Exception as e:
        print(f"Error converting markdown to DOCX: {e}", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
